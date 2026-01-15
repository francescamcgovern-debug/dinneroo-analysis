#!/usr/bin/env python3
"""
Generate Word document for User ID Linkage Guide
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

# Output path
OUTPUT_PATH = Path(__file__).parent.parent / "DELIVERABLES" / "reports" / "User_ID_Linkage_Guide.docx"

def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    
    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)
    
    doc.add_paragraph()  # spacing

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('User ID to Survey Response Linkage Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Purpose
    add_heading(doc, 'Purpose', 1)
    doc.add_paragraph(
        'This guide documents the process used in the Dinneroo project to connect customer identities '
        '(User IDs and emails) from Snowflake order data to survey responses collected via Alchemer. '
        'This enables us to enrich survey feedback with actual behavioral data (order history, ratings, '
        'zone, dishes ordered).'
    )
    
    # Data Flow Architecture
    add_heading(doc, 'Data Flow Architecture', 1)
    doc.add_paragraph(
        'The linkage process follows this flow:\n\n'
        '1. Snowflake Orders → Pull Customer Data (with emails)\n'
        '2. Alchemer Surveys → Consolidate Survey Responses\n'
        '3. Both sources → Enrich via Linkage\n'
        '4. Output: post_order_enriched_COMPLETE.csv and DROPOFF_ENRICHED.csv'
    )
    
    # Step 1
    add_heading(doc, 'Step 1: Pull Customer Data from Snowflake', 1)
    
    add_heading(doc, 'Script', 2)
    doc.add_paragraph('integrations/snowflake_connector.py')
    
    add_heading(doc, 'How It Works', 2)
    doc.add_paragraph(
        'The Snowflake connector pulls Dinneroo order data with customer identifiers by joining '
        'the orders table with the users table. The key SQL logic is:'
    )
    doc.add_paragraph(
        'SELECT o.USER_ID AS CUSTOMER_ID, u.EMAIL AS CUSTOMER_EMAIL, ...\n'
        'FROM production.denormalised.orders o\n'
        'LEFT JOIN production.orderweb.users u ON o.USER_ID = u.ID\n'
        'WHERE o.IS_FAMILY_TIME_ORDER = TRUE AND o.STATUS = \'DELIVERED\''
    )
    
    add_heading(doc, 'Key Columns for Linkage', 2)
    add_table(doc, 
        ['Column', 'Purpose'],
        [
            ['CUSTOMER_ID', 'Internal Deliveroo User ID'],
            ['CUSTOMER_EMAIL', 'Email for matching to surveys'],
            ['ORDER_ID', 'Unique order identifier'],
            ['ORDER_TIMESTAMP', 'For date-proximity matching'],
            ['PARTNER_NAME', 'For brand-based fallback matching'],
        ]
    )
    
    add_heading(doc, 'Output', 2)
    doc.add_paragraph('DATA/1_SOURCE/snowflake/ALL_DINNEROO_ORDERS.csv')
    
    add_heading(doc, 'Running the Data Pull', 2)
    doc.add_paragraph('python scripts/refresh_snowflake_data.py')
    
    # Step 2
    add_heading(doc, 'Step 2: Consolidate Survey Responses', 1)
    
    add_heading(doc, 'Script', 2)
    doc.add_paragraph('scripts/phase1_data/04_consolidate_surveys.py')
    
    add_heading(doc, 'How It Works', 2)
    doc.add_paragraph(
        'Surveys are exported from Alchemer as CSV files. This script:\n\n'
        '1. Loads existing consolidated survey file (if any)\n'
        '2. Loads new Alchemer export\n'
        '3. Merges and deduplicates on Response ID\n'
        '4. Saves updated consolidated file'
    )
    
    add_heading(doc, 'Survey Email Field', 2)
    doc.add_paragraph(
        'The survey asks respondents for their email with this question:\n'
        '"Please provide your Deliveroo email address so that we can send you the voucher!"\n\n'
        'This email is the primary linkage key to Snowflake data.'
    )
    
    add_heading(doc, 'Running Consolidation', 2)
    doc.add_paragraph('python scripts/phase1_data/04_consolidate_surveys.py --all')
    
    # Step 3
    add_heading(doc, 'Step 3: Link Survey Responses to Order Data', 1)
    
    add_heading(doc, 'Script', 2)
    doc.add_paragraph('scripts/phase1_data/05_enrich_surveys.py')
    
    add_heading(doc, 'Linkage Strategy (Confidence Hierarchy)', 2)
    doc.add_paragraph(
        'The script attempts to match each survey response to a Snowflake order using a tiered confidence approach:'
    )
    
    add_table(doc,
        ['Priority', 'Method', 'Confidence', 'Logic'],
        [
            ['1', 'email_brand_date', 'Highest', 'Email + restaurant brand + order 0-7 days before survey'],
            ['2', 'email_date', 'High', 'Email + order 0-7 days before survey'],
            ['3', 'email_only', 'Medium', 'Email matches, use most recent order'],
            ['4', 'brand_date', 'Low', 'Restaurant brand + date match (may match wrong customer)'],
            ['5', 'unmatched', 'None', 'No linkage possible'],
        ]
    )
    
    add_heading(doc, 'Key Functions', 2)
    doc.add_paragraph(
        '1. normalize_email() - Lowercase and strip whitespace for matching\n'
        '2. normalize_brand_name() - Extract brand from partner name (e.g., "Pho - Bristol" → "pho")\n'
        '3. build_email_to_orders_index() - Create lookup dict: {email: [order_rows]}\n'
        '4. find_best_order_match() - Apply linkage hierarchy'
    )
    
    add_heading(doc, 'Date Matching Logic', 2)
    doc.add_paragraph(
        'Orders are considered valid matches if they occurred 0-7 days before the survey submission. '
        'The most recent matching order (smallest days_diff) is preferred.'
    )
    
    add_heading(doc, 'Enriched Fields Added', 2)
    add_table(doc,
        ['Field', 'Source', 'Description'],
        [
            ['ORDER_ID', 'Matched order', 'Unique order identifier'],
            ['USER_ID', 'Orders table', 'Deliveroo Customer ID'],
            ['USER_EMAIL', 'Orders table', 'Customer email'],
            ['LINKAGE_METHOD', 'Enrichment', 'Which strategy succeeded'],
            ['ORDER_TIMESTAMP', 'Orders table', 'When order was placed'],
            ['PARTNER_NAME', 'Orders table', 'Restaurant name'],
            ['ORDER_VALUE', 'Orders table', 'Order total'],
            ['ZONE_NAME', 'Orders table', 'Delivery zone'],
            ['MENU_ITEM_LIST', 'Orders table', 'Items ordered'],
            ['RATING_STARS', 'Ratings table', 'Customer rating (1-5)'],
            ['RATING_COMMENT', 'Ratings table', 'Rating text feedback'],
        ]
    )
    
    add_heading(doc, 'Derived Metrics', 2)
    add_table(doc,
        ['Metric', 'Logic'],
        [
            ['NUM_CHILDREN_NUMERIC', 'Count of child age columns with values'],
            ['Is_Family', 'Has children OR ordered for family'],
            ['Is_Satisfied', 'Rating is "Very Satisfied" or "Satisfied"'],
            ['Days_To_Survey', 'Days between order and survey submission'],
            ['Strong_Advocate', 'Would reorder + satisfied'],
        ]
    )
    
    add_heading(doc, 'Running Enrichment', 2)
    doc.add_paragraph('python scripts/phase1_data/05_enrich_surveys.py --all')
    
    # Step 4
    add_heading(doc, 'Step 4: Complete Pipeline', 1)
    
    add_heading(doc, 'Full Refresh Command Sequence', 2)
    doc.add_paragraph(
        '# 1. Pull fresh data from Snowflake (requires SSO)\n'
        'python scripts/refresh_snowflake_data.py\n\n'
        '# 2. Consolidate survey exports from Alchemer\n'
        'python scripts/phase1_data/04_consolidate_surveys.py --all\n\n'
        '# 3. Enrich surveys with Snowflake linkage\n'
        'python scripts/phase1_data/05_enrich_surveys.py --all'
    )
    
    # Match Rates
    add_heading(doc, 'Match Rate Analysis', 1)
    
    add_heading(doc, 'Typical Linkage Breakdown', 2)
    add_table(doc,
        ['Linkage Method', 'Typical %', 'Notes'],
        [
            ['email_brand_date', '40-50%', 'Best quality matches'],
            ['email_date', '20-30%', 'Good quality, brand mismatch'],
            ['email_only', '5-10%', 'Older orders, timing gap'],
            ['brand_date', '5-10%', 'Fallback, lower confidence'],
            ['unmatched', '10-20%', 'No email or not in Snowflake'],
        ]
    )
    
    add_heading(doc, 'Why Matches Fail', 2)
    doc.add_paragraph(
        '1. No email provided - Respondent skipped email field\n'
        '2. Different email - Used different email for survey vs Deliveroo account\n'
        '3. New customer - Placed first order, not yet in historical data\n'
        '4. Timing gap - Order was >7 days before survey\n'
        '5. Data lag - Snowflake data not yet refreshed'
    )
    
    # Key Files
    add_heading(doc, 'Key Files Reference', 1)
    add_table(doc,
        ['Purpose', 'Path'],
        [
            ['Snowflake connector', 'integrations/snowflake_connector.py'],
            ['Data refresh script', 'scripts/refresh_snowflake_data.py'],
            ['Survey consolidation', 'scripts/phase1_data/04_consolidate_surveys.py'],
            ['Survey enrichment', 'scripts/phase1_data/05_enrich_surveys.py'],
            ['Raw orders', 'DATA/1_SOURCE/snowflake/ALL_DINNEROO_ORDERS.csv'],
            ['Raw customers', 'DATA/1_SOURCE/snowflake/ALL_DINNEROO_CUSTOMERS.csv'],
            ['Enriched post-order', 'DATA/2_ENRICHED/post_order_enriched_COMPLETE.csv'],
            ['Enriched dropoff', 'DATA/2_ENRICHED/DROPOFF_ENRICHED.csv'],
        ]
    )
    
    # Considerations
    add_heading(doc, 'Important Considerations', 1)
    
    add_heading(doc, 'Data Quality', 2)
    doc.add_paragraph(
        '1. Email is the primary linkage key - Survey respondents must provide their Deliveroo email\n'
        '2. Date proximity matters - Orders 0-7 days before survey are considered\n'
        '3. Brand validation increases confidence - If restaurant in survey matches order, higher trust\n'
        '4. Unmatched responses are preserved - Survey data retained even without order linkage'
    )
    
    add_heading(doc, 'Privacy & Compliance', 2)
    doc.add_paragraph(
        '1. GDPR compliance - Email data only used for internal analysis linkage\n'
        '2. Data minimization - Only necessary fields are pulled from Snowflake\n'
        '3. Access control - Snowflake access requires authenticated credentials\n'
        '4. No PII in outputs - Enriched files should not be shared externally with raw emails'
    )
    
    add_heading(doc, 'Troubleshooting', 2)
    add_table(doc,
        ['Issue', 'Likely Cause', 'Solution'],
        [
            ['Low match rate', 'Stale Snowflake data', 'Run refresh_snowflake_data.py'],
            ['Missing CUSTOMER_EMAIL', 'Query not joining users table', 'Check connector query'],
            ['All brand_date matches', 'Email column not found', 'Check survey email field name'],
            ['Zero matches', 'Wrong file paths', 'Verify DATA/1_SOURCE/ structure'],
        ]
    )
    
    # Save
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Word document saved to: {OUTPUT_PATH}")

if __name__ == "__main__":
    main()
